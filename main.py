from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse
import tempfile
import os
import magic
from typing import Dict, Any
import fitz  # PyMuPDF for PDF processing
from PIL import Image
import openpyxl
from docx import Document
import pandas as pd

app = FastAPI(title="Document Metadata Extractor", version="1.0.0")

SUPPORTED_TYPES = {
    'application/pdf': ['.pdf'],
    'application/vnd.ms-excel': ['.xls'],
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': ['.xlsx'],
    'image/jpeg': ['.jpg', '.jpeg'],
    'image/png': ['.png'],
    'image/tiff': ['.tif', '.tiff'],
    'text/plain': ['.txt'],
    'application/msword': ['.doc'],
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx']
}

def process_pdf(file_path: str) -> Dict[str, Any]:
    """Process PDF file and extract metadata"""
    doc = fitz.open(file_path)
    metadata = {
        "document_type": "pdf",
        "page_count": len(doc),
        "text_preview": "",
        "tables_count": 0,
        "images_count": 0,
        "structure_elements": ["PDF"],
        "has_bookmarks": len(doc.get_toc()) > 0,
        "is_encrypted": doc.is_encrypted
    }
    
    # Extract text from first few pages
    text_content = ""
    for page_num in range(min(3, len(doc))):
        page = doc[page_num]
        text_content += page.get_text()
        
        # Count images on page
        image_list = page.get_images()
        metadata["images_count"] += len(image_list)
    
    metadata["text_preview"] = text_content[:500] + "..." if len(text_content) > 500 else text_content
    doc.close()
    return metadata

def process_excel(file_path: str) -> Dict[str, Any]:
    """Process Excel file and extract metadata"""
    sheet_names = []
    text_content = ""
    total_rows = 0
    
    try:
        # Use pandas to read all sheets and extract content
        excel_file = pd.ExcelFile(file_path)
        sheet_names = excel_file.sheet_names
        
        # Extract content from each sheet (limit to first 3 sheets and first 100 rows per sheet)
        for sheet_name in sheet_names[:3]:
            try:
                df = pd.read_excel(file_path, sheet_name=sheet_name, nrows=100)
                total_rows += len(df)
                
                # Add sheet header
                text_content += f"\n=== PLANILHA: {sheet_name} ===\n"
                
                # Add column headers
                text_content += "COLUNAS: " + " | ".join(df.columns.astype(str)) + "\n\n"
                
                # Add first few rows of data
                for idx, row in df.head(20).iterrows():
                    row_text = " | ".join([str(val) if pd.notna(val) else "" for val in row])
                    text_content += f"Linha {idx + 1}: {row_text}\n"
                
                text_content += f"\n[{len(df)} linhas nesta planilha]\n"
                
            except Exception as e:
                text_content += f"\n=== PLANILHA: {sheet_name} ===\n"
                text_content += f"Erro ao ler planilha: {str(e)}\n"
        
        excel_file.close()
        
    except Exception as e:
        # Fallback method using openpyxl
        try:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            sheet_names = wb.sheetnames
            
            for sheet_name in sheet_names[:2]:  # Limit to first 2 sheets
                ws = wb[sheet_name]
                text_content += f"\n=== PLANILHA: {sheet_name} ===\n"
                
                # Get data from first 50 rows and 20 columns
                rows_data = []
                for row in ws.iter_rows(max_row=50, max_col=20, values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():  # Only add non-empty rows
                        rows_data.append(row_text)
                
                text_content += "\n".join(rows_data[:30])  # First 30 non-empty rows
                text_content += f"\n\n[Planilha {sheet_name} processada]\n"
            
            wb.close()
            
        except Exception as e2:
            text_content = f"Erro ao processar Excel: {str(e)} / {str(e2)}"
            sheet_names = ["Erro"]
    
    # Create preview (first 1500 characters)
    preview = text_content[:1500] + "..." if len(text_content) > 1500 else text_content
    
    metadata = {
        "document_type": "spreadsheet",
        "page_count": len(sheet_names),
        "sheet_names": sheet_names,
        "text_preview": preview,
        "full_content": text_content,  # Full content for prompt use
        "total_rows": total_rows,
        "tables_count": len(sheet_names),
        "images_count": 0,
        "structure_elements": ["Spreadsheet", "Worksheets", "Data"]
    }
    return metadata

def process_image(file_path: str) -> Dict[str, Any]:
    """Process image file and extract metadata"""
    try:
        with Image.open(file_path) as img:
            metadata = {
                "document_type": "image",
                "page_count": 1,
                "dimensions": f"{img.width}x{img.height}",
                "format": img.format,
                "mode": img.mode,
                "text_preview": f"Image file ({img.format}, {img.width}x{img.height}, {img.mode})",
                "tables_count": 0,
                "images_count": 1,
                "structure_elements": ["Image"]
            }
            return metadata
    except Exception as e:
        return {
            "document_type": "image",
            "page_count": 1,
            "text_preview": "Image file (unable to read metadata)",
            "tables_count": 0,
            "images_count": 1,
            "structure_elements": ["Image"],
            "error": str(e)
        }

def process_docx(file_path: str) -> Dict[str, Any]:
    """Process Word document and extract metadata"""
    try:
        doc = Document(file_path)
        text_content = ""
        table_count = 0
        
        for paragraph in doc.paragraphs[:10]:  # First 10 paragraphs
            text_content += paragraph.text + "\n"
        
        table_count = len(doc.tables)
        
        metadata = {
            "document_type": "document",
            "page_count": len(doc.sections),
            "paragraph_count": len(doc.paragraphs),
            "text_preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
            "tables_count": table_count,
            "images_count": 0,
            "structure_elements": ["Document", "Paragraphs", "Tables"] if table_count > 0 else ["Document", "Paragraphs"]
        }
        return metadata
    except Exception as e:
        return {
            "document_type": "document",
            "page_count": 1,
            "text_preview": "Word document (unable to read content)",
            "tables_count": 0,
            "images_count": 0,
            "structure_elements": ["Document"],
            "error": str(e)
        }

def process_text(file_path: str) -> Dict[str, Any]:
    """Process text file and extract metadata"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        metadata = {
            "document_type": "text",
            "page_count": 1,
            "line_count": len(lines),
            "character_count": len(content),
            "text_preview": content[:500] + "..." if len(content) > 500 else content,
            "tables_count": 0,
            "images_count": 0,
            "structure_elements": ["Text"]
        }
        return metadata
    except Exception as e:
        return {
            "document_type": "text",
            "page_count": 1,
            "text_preview": "Text file (unable to read content)",
            "tables_count": 0,
            "images_count": 0,
            "structure_elements": ["Text"],
            "error": str(e)
        }

def extract_metadata_for_prompt(file_path: str, mime_type: str) -> Dict[str, Any]:
    """Extract relevant metadata from document for prompt context"""
    
    if mime_type == 'application/pdf':
        return process_pdf(file_path)
    elif mime_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
        return process_excel(file_path)
    elif mime_type.startswith('image/'):
        return process_image(file_path)
    elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        return process_docx(file_path)
    elif mime_type == 'text/plain':
        return process_text(file_path)
    else:
        return {
            "document_type": "unknown",
            "page_count": 1,
            "text_preview": "Unsupported document type",
            "tables_count": 0,
            "images_count": 0,
            "structure_elements": ["Unknown"]
        }

@app.post("/extract-metadata")
async def extract_document_metadata(file: UploadFile = File(...)):
    """
    Extract metadata from uploaded document (PDF, XLS, JPG, etc.) for prompt generation
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as temp_file:
        content = await file.read()
        temp_file.write(content)
        temp_file_path = temp_file.name
    
    try:
        # Detect file type
        mime_type = magic.from_file(temp_file_path, mime=True)
        
        # Validate supported file type
        if mime_type not in SUPPORTED_TYPES:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type: {mime_type}. Supported types: {list(SUPPORTED_TYPES.keys())}"
            )
        
        # Extract metadata
        metadata = extract_metadata_for_prompt(temp_file_path, mime_type)
        
        # Add file info
        metadata.update({
            "filename": file.filename,
            "file_size": len(content),
            "mime_type": mime_type,
            "file_extension": os.path.splitext(file.filename)[1]
        })
        
        # Create enhanced prompt context
        prompt_context = f"""Document Analysis:
- Type: {metadata['document_type']}
- Pages: {metadata.get('page_count', 1)}
- Tables: {metadata.get('tables_count', 0)}
- Images: {metadata.get('images_count', 0)}
- Structure: {', '.join(metadata.get('structure_elements', ['Unknown']))}

Content:
{metadata.get('full_content') if 'full_content' in metadata else metadata.get('text_preview', 'No content available')}"""

        return JSONResponse(content={
            "success": True,
            "metadata": metadata,
            "prompt_context": prompt_context
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")
    
    finally:
        # Clean up temporary file
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)

@app.get("/")
async def root():
    return {"message": "Document Metadata Extractor API", "endpoints": ["/extract-metadata"]}

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)